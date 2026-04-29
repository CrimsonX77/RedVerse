#!/usr/bin/env python3
"""
Dragon Forge Server — Format-agnostic conversion backend
Drop any file. Pick any compatible target. Forge it.

Port 8916 | Part of RedVerse Agency Scripts

Core deps (required):
    pip install flask flask-cors Pillow

Recommended deps (each unlocks a conversion lane):
    pip install PyMuPDF python-docx markdown PyYAML html2text reportlab cairosvg

System binaries:
    ffmpeg         → audio/video conversion (essential)
    potrace        → raster → true-vector SVG (optional)
"""
import sys
import os
import io
import re
import json
import csv
import tempfile
import shutil
import subprocess
from pathlib import Path
from dataclasses import dataclass
from typing import Callable, Set, Optional
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024 * 1024  # 2GB

# ═══════════════════════════════════════════════════════════════
# DEPENDENCY PROBE
# ═══════════════════════════════════════════════════════════════

DEPS = {
    'pil': False, 'ffmpeg': False, 'pymupdf': False, 'docx': False,
    'markdown': False, 'yaml': False, 'html2text': False,
    'reportlab': False, 'cairosvg': False, 'potrace': False,
}

def probe_deps():
    try:
        import PIL  # noqa
        DEPS['pil'] = True
    except ImportError: pass
    DEPS['ffmpeg']  = shutil.which('ffmpeg') is not None
    DEPS['potrace'] = shutil.which('potrace') is not None
    for mod, key in [('fitz','pymupdf'),('docx','docx'),('markdown','markdown'),
                     ('yaml','yaml'),('html2text','html2text'),
                     ('reportlab','reportlab'),('cairosvg','cairosvg')]:
        try:
            __import__(mod); DEPS[key] = True
        except ImportError: pass

probe_deps()

# ═══════════════════════════════════════════════════════════════
# FORMAT REGISTRY
# ═══════════════════════════════════════════════════════════════

CANONICAL = {
    'jpeg':'jpg', 'htm':'html', 'yml':'yaml', 'tif':'tiff', 'aif':'aiff',
    'mpeg':'mpg',
}

def canon(ext: str) -> str:
    ext = ext.lower().lstrip('.')
    return CANONICAL.get(ext, ext)

FORMATS = {
    # raster images
    'png':  {'cat':'image',    'desc':'PNG',    'vector':False},
    'jpg':  {'cat':'image',    'desc':'JPEG',   'vector':False},
    'webp': {'cat':'image',    'desc':'WebP',   'vector':False},
    'bmp':  {'cat':'image',    'desc':'Bitmap', 'vector':False},
    'tiff': {'cat':'image',    'desc':'TIFF',   'vector':False},
    'gif':  {'cat':'image',    'desc':'GIF',    'vector':False},
    'ico':  {'cat':'image',    'desc':'Icon',   'vector':False},
    # vector images
    'svg':  {'cat':'image',    'desc':'SVG',    'vector':True},
    # audio
    'mp3':  {'cat':'audio',    'desc':'MP3'},
    'wav':  {'cat':'audio',    'desc':'WAV'},
    'flac': {'cat':'audio',    'desc':'FLAC'},
    'aac':  {'cat':'audio',    'desc':'AAC'},
    'ogg':  {'cat':'audio',    'desc':'OGG'},
    'opus': {'cat':'audio',    'desc':'Opus'},
    'm4a':  {'cat':'audio',    'desc':'M4A'},
    'aiff': {'cat':'audio',    'desc':'AIFF'},
    # video
    'mp4':  {'cat':'video',    'desc':'MP4'},
    'mkv':  {'cat':'video',    'desc':'MKV'},
    'webm': {'cat':'video',    'desc':'WebM'},
    'avi':  {'cat':'video',    'desc':'AVI'},
    'mov':  {'cat':'video',    'desc':'MOV'},
    # text
    'txt':  {'cat':'text',     'desc':'Plain Text'},
    'md':   {'cat':'text',     'desc':'Markdown'},
    'html': {'cat':'text',     'desc':'HTML'},
    'json': {'cat':'text',     'desc':'JSON'},
    'yaml': {'cat':'text',     'desc':'YAML'},
    'xml':  {'cat':'text',     'desc':'XML'},
    'csv':  {'cat':'text',     'desc':'CSV'},
    # documents
    'pdf':  {'cat':'document', 'desc':'PDF'},
    'docx': {'cat':'document', 'desc':'Word Doc'},
}

# ═══════════════════════════════════════════════════════════════
# CONVERTER HANDLERS
# ═══════════════════════════════════════════════════════════════

@dataclass
class Rule:
    name: str
    sources: Set[str]
    targets: Set[str]
    check: Callable[[], bool]
    handler: Callable
    note: str = ''

# ─── Raster / vector images ────────────────────────────────────

def H_raster_raster(in_p, out_p, src, tgt, q):
    from PIL import Image
    img = Image.open(in_p)
    ext = '.' + tgt
    if ext in ('.jpg','.bmp','.ico') and img.mode == 'RGBA':
        bg = Image.new('RGB', img.size, (0,0,0))
        bg.paste(img, mask=img.split()[3]); img = bg
    elif ext in ('.jpg','.bmp') and img.mode != 'RGB':
        img = img.convert('RGB')
    kw = {}
    if ext == '.jpg':   kw.update(quality=q, optimize=True)
    elif ext == '.webp': kw.update(quality=q)
    elif ext == '.png':  kw.update(optimize=True)
    img.save(out_p, **kw)

def H_svg_raster(in_p, out_p, src, tgt, q):
    import cairosvg
    if tgt == 'png':
        cairosvg.svg2png(url=in_p, write_to=out_p)
    elif tgt == 'pdf':
        cairosvg.svg2pdf(url=in_p, write_to=out_p)
    else:
        tmp = out_p + '.tmp.png'
        cairosvg.svg2png(url=in_p, write_to=tmp)
        H_raster_raster(tmp, out_p, 'png', tgt, q)
        os.remove(tmp)

def H_raster_svg(in_p, out_p, src, tgt, q):
    with tempfile.NamedTemporaryFile(suffix='.bmp', delete=False) as tmp:
        bmp_path = tmp.name
    try:
        from PIL import Image
        Image.open(in_p).convert('L').save(bmp_path, 'BMP')
        subprocess.run(['potrace', bmp_path, '-s', '-o', out_p], check=True,
                       capture_output=True)
    finally:
        if os.path.exists(bmp_path): os.remove(bmp_path)

def H_raster_pdf(in_p, out_p, src, tgt, q):
    from PIL import Image
    img = Image.open(in_p)
    if img.mode != 'RGB': img = img.convert('RGB')
    img.save(out_p, 'PDF')

# ─── Audio / video via ffmpeg ──────────────────────────────────

def H_ffmpeg(in_p, out_p, src, tgt, q):
    cmd = ['ffmpeg','-y','-i',in_p]
    ext = '.' + tgt
    src_cat = FORMATS[src]['cat']
    tgt_cat = FORMATS[tgt]['cat']

    if src_cat == 'video' and tgt_cat == 'audio':
        cmd.extend(['-vn'])
    if src_cat == 'video' and tgt_cat == 'image':
        cmd.extend(['-ss','00:00:01','-frames:v','1'])

    if tgt_cat == 'audio':
        if ext == '.mp3':    cmd.extend(['-b:a', f'{int(64+(q/100)*256)}k'])
        elif ext == '.flac': cmd.extend(['-compression_level','8'])
        elif ext == '.ogg':  cmd.extend(['-q:a', str(max(0,min(10,int(q/10))))])
        elif ext == '.aac':  cmd.extend(['-b:a', f'{int(64+(q/100)*192)}k'])
        elif ext == '.opus': cmd.extend(['-b:a', f'{int(32+(q/100)*224)}k'])
        elif ext == '.m4a':  cmd.extend(['-c:a','aac','-b:a', f'{int(64+(q/100)*192)}k'])
    elif tgt_cat == 'video':
        if ext == '.gif':
            cmd.extend(['-vf','fps=15,scale=480:-1:flags=lanczos','-loop','0'])
        elif ext in ('.mp4','.mov'):
            crf = int(51-(q/100)*41)
            cmd.extend(['-c:v','libx264','-crf',str(crf),'-preset','medium'])
            if ext == '.mp4': cmd.extend(['-movflags','+faststart'])
        elif ext == '.webm':
            crf = int(63-(q/100)*53)
            cmd.extend(['-c:v','libvpx-vp9','-crf',str(crf),'-b:v','0'])
        elif ext == '.mkv':
            cmd.extend(['-c:v','libx264','-crf',str(int(51-(q/100)*41))])
        elif ext == '.avi':
            cmd.extend(['-c:v','mpeg4','-q:v', str(max(2,int(31-(q/100)*29)))])

    cmd.append(out_p)
    proc = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    _, stderr = proc.communicate()
    if proc.returncode != 0:
        raise RuntimeError(f'ffmpeg: {stderr.decode(errors="ignore")[-400:]}')

# ─── Text transforms ───────────────────────────────────────────

def _read(path):
    with open(path,'r',encoding='utf-8',errors='replace') as f: return f.read()
def _write(path,s):
    with open(path,'w',encoding='utf-8') as f: f.write(s)

def _yaml_loads(s):
    import yaml; return yaml.safe_load(s)
def _yaml_dumps(o):
    import yaml; return yaml.safe_dump(o, sort_keys=False, allow_unicode=True)

def _xml_to_dict(s):
    import xml.etree.ElementTree as ET
    root = ET.fromstring(s)
    def walk(el):
        kids = list(el)
        if not kids: return (el.text or '').strip()
        d = {}
        for c in kids:
            v = walk(c)
            if c.tag in d:
                if not isinstance(d[c.tag], list): d[c.tag] = [d[c.tag]]
                d[c.tag].append(v)
            else: d[c.tag] = v
        return d
    return {root.tag: walk(root)}

def _text_transform(data, src, tgt):
    src, tgt = canon(src), canon(tgt)
    if src == tgt: return data

    # JSON ↔ YAML
    if src in ('json','yaml') and tgt in ('json','yaml'):
        obj = json.loads(data) if src == 'json' else _yaml_loads(data)
        return json.dumps(obj, indent=2) if tgt == 'json' else _yaml_dumps(obj)

    # MD → HTML
    if src == 'md' and tgt == 'html':
        import markdown
        body = markdown.markdown(data, extensions=['fenced_code','tables'])
        return f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{body}</body></html>"
    # HTML → MD
    if src == 'html' and tgt == 'md':
        import html2text
        h = html2text.HTML2Text(); h.body_width = 0
        return h.handle(data)
    # HTML → TXT
    if src == 'html' and tgt == 'txt':
        t = re.sub(r'<script[^>]*>.*?</script>','',data,flags=re.DOTALL|re.I)
        t = re.sub(r'<style[^>]*>.*?</style>','',t,flags=re.DOTALL|re.I)
        t = re.sub(r'<[^>]+>','',t)
        return re.sub(r'\s+\n','\n',t).strip()
    # MD → TXT
    if src == 'md' and tgt == 'txt':
        t = re.sub(r'^#+\s*','',data,flags=re.M)
        t = re.sub(r'\*\*([^*]+)\*\*',r'\1',t)
        t = re.sub(r'\*([^*]+)\*',r'\1',t)
        t = re.sub(r'`([^`]+)`',r'\1',t)
        return re.sub(r'\[([^\]]+)\]\([^)]+\)',r'\1',t)
    # TXT → MD/HTML
    if src == 'txt' and tgt == 'md': return data
    if src == 'txt' and tgt == 'html':
        esc = data.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        return f"<!DOCTYPE html><html><body><pre>{esc}</pre></body></html>"

    # CSV ↔ JSON
    if src == 'csv' and tgt == 'json':
        return json.dumps(list(csv.DictReader(io.StringIO(data))), indent=2)
    if src == 'json' and tgt == 'csv':
        obj = json.loads(data)
        if not isinstance(obj, list) or not obj: return ''
        keys = list(obj[0].keys())
        buf = io.StringIO()
        w = csv.DictWriter(buf, fieldnames=keys); w.writeheader()
        for row in obj: w.writerow({k: row.get(k,'') for k in keys})
        return buf.getvalue()

    # CSV ↔ YAML
    if src == 'csv' and tgt == 'yaml':
        return _yaml_dumps(list(csv.DictReader(io.StringIO(data))))
    if src == 'yaml' and tgt == 'csv':
        obj = _yaml_loads(data)
        if not isinstance(obj, list) or not obj: return ''
        keys = list(obj[0].keys())
        buf = io.StringIO()
        w = csv.DictWriter(buf, fieldnames=keys); w.writeheader()
        for row in obj: w.writerow({k: row.get(k,'') for k in keys})
        return buf.getvalue()

    # XML → JSON/YAML
    if src == 'xml' and tgt == 'json': return json.dumps(_xml_to_dict(data), indent=2)
    if src == 'xml' and tgt == 'yaml': return _yaml_dumps(_xml_to_dict(data))

    # Compound paths: anything → intermediate → target
    if src == 'md' and tgt in ('json','yaml','csv','xml'):
        # MD is prose, best we can do is wrap as a list of lines
        return _text_transform(_text_transform(data,'md','txt'),'txt',tgt)
    if src == 'html' and tgt in ('json','yaml','csv','xml'):
        return _text_transform(_text_transform(data,'html','txt'),'txt',tgt)
    if src == 'txt' and tgt in ('json','yaml'):
        # Wrap plain text as a string list
        obj = data.splitlines()
        return json.dumps(obj, indent=2) if tgt == 'json' else _yaml_dumps(obj)

    return data  # fallback — pass through

def H_text_text(in_p, out_p, src, tgt, q):
    _write(out_p, _text_transform(_read(in_p), src, tgt))

# ─── Documents ─────────────────────────────────────────────────

def H_text_pdf(in_p, out_p, src, tgt, q):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm
    data = _read(in_p)
    if src == 'html': data = _text_transform(data,'html','txt')
    elif src == 'md': data = _text_transform(data,'md','txt')
    c = canvas.Canvas(out_p, pagesize=A4)
    _, h = A4
    x, y = 2*cm, h - 2*cm
    c.setFont('Helvetica', 10)
    for line in data.splitlines():
        while len(line) > 100:
            c.drawString(x, y, line[:100]); line = line[100:]; y -= 12
            if y < 2*cm: c.showPage(); c.setFont('Helvetica',10); y = h - 2*cm
        c.drawString(x, y, line); y -= 12
        if y < 2*cm: c.showPage(); c.setFont('Helvetica',10); y = h - 2*cm
    c.save()

def H_text_docx(in_p, out_p, src, tgt, q):
    import docx
    data = _read(in_p)
    if src == 'html': data = _text_transform(data,'html','txt')
    elif src == 'md': data = _text_transform(data,'md','txt')
    d = docx.Document()
    for line in data.splitlines(): d.add_paragraph(line)
    d.save(out_p)

def H_pdf_text(in_p, out_p, src, tgt, q):
    import fitz
    doc = fitz.open(in_p)
    joined = '\n\n'.join(p.get_text() for p in doc); doc.close()
    if tgt == 'txt': _write(out_p, joined)
    else: _write(out_p, _text_transform(joined,'txt',tgt))

def H_pdf_image(in_p, out_p, src, tgt, q):
    import fitz
    doc = fitz.open(in_p)
    page = doc[0]
    pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
    tmp = out_p + '.tmp.png'
    pix.save(tmp); doc.close()
    if tgt == 'png': shutil.move(tmp, out_p)
    else:
        H_raster_raster(tmp, out_p, 'png', tgt, q)
        os.remove(tmp)

def H_docx_text(in_p, out_p, src, tgt, q):
    import docx
    d = docx.Document(in_p)
    joined = '\n'.join(p.text for p in d.paragraphs)
    if tgt == 'txt': _write(out_p, joined)
    else: _write(out_p, _text_transform(joined,'txt',tgt))

def H_pdf_docx(in_p, out_p, src, tgt, q):
    import fitz, docx
    d = fitz.open(in_p)
    text = '\n\n'.join(p.get_text() for p in d); d.close()
    out = docx.Document()
    for line in text.splitlines(): out.add_paragraph(line)
    out.save(out_p)

def H_docx_pdf(in_p, out_p, src, tgt, q):
    import docx
    d = docx.Document(in_p)
    text = '\n'.join(p.text for p in d.paragraphs)
    tmp = tempfile.NamedTemporaryFile(suffix='.txt', delete=False, mode='w', encoding='utf-8')
    tmp.write(text); tmp.close()
    try: H_text_pdf(tmp.name, out_p, 'txt', 'pdf', q)
    finally: os.remove(tmp.name)

# ─── Rule table ────────────────────────────────────────────────

RASTER = {'png','jpg','webp','bmp','tiff','gif','ico'}
AUDIO  = {'mp3','wav','flac','aac','ogg','opus','m4a','aiff'}
VIDEO  = {'mp4','mkv','webm','avi','mov'}
TEXT   = {'txt','md','html','json','yaml','xml','csv'}

RULES = [
    Rule('raster↔raster',  RASTER,           RASTER,                lambda: DEPS['pil'],                            H_raster_raster),
    Rule('raster→pdf',     RASTER,           {'pdf'},               lambda: DEPS['pil'],                            H_raster_pdf),
    Rule('svg→raster',     {'svg'},          RASTER | {'pdf'},      lambda: DEPS['cairosvg'],                       H_svg_raster,     'cairosvg required'),
    Rule('raster→svg',     RASTER,           {'svg'},               lambda: DEPS['pil'] and DEPS['potrace'],        H_raster_svg,     'potrace binary required'),
    Rule('audio↔audio',    AUDIO,            AUDIO,                 lambda: DEPS['ffmpeg'],                         H_ffmpeg),
    Rule('video↔video',    VIDEO | {'gif'},  VIDEO | {'gif'},       lambda: DEPS['ffmpeg'],                         H_ffmpeg),
    Rule('video→audio',    VIDEO,            AUDIO,                 lambda: DEPS['ffmpeg'],                         H_ffmpeg),
    Rule('video→image',    VIDEO,            RASTER,                lambda: DEPS['ffmpeg'],                         H_ffmpeg),
    Rule('text↔text',      TEXT,             TEXT,                  lambda: True,                                   H_text_text,      'some lanes need markdown/html2text/PyYAML'),
    Rule('text→pdf',       TEXT,             {'pdf'},               lambda: DEPS['reportlab'],                      H_text_pdf),
    Rule('text→docx',      TEXT,             {'docx'},              lambda: DEPS['docx'],                           H_text_docx),
    Rule('pdf→text',       {'pdf'},          TEXT,                  lambda: DEPS['pymupdf'],                        H_pdf_text),
    Rule('pdf→image',      {'pdf'},          RASTER,                lambda: DEPS['pymupdf'] and DEPS['pil'],        H_pdf_image),
    Rule('pdf→docx',       {'pdf'},          {'docx'},              lambda: DEPS['pymupdf'] and DEPS['docx'],       H_pdf_docx),
    Rule('docx→text',      {'docx'},         TEXT,                  lambda: DEPS['docx'],                           H_docx_text),
    Rule('docx→pdf',       {'docx'},         {'pdf'},               lambda: DEPS['docx'] and DEPS['reportlab'],     H_docx_pdf),
]

def valid_targets_for(src_ext: str):
    src = canon(src_ext)
    if src not in FORMATS: return []
    targets = {}
    for rule in RULES:
        if src not in rule.sources or not rule.check(): continue
        for t in rule.targets:
            if t == src or t not in FORMATS: continue
            if t not in targets:
                targets[t] = {'ext':t,'desc':FORMATS[t]['desc'],
                              'cat':FORMATS[t]['cat'],'via':rule.name}
    return sorted(targets.values(), key=lambda x: (x['cat'], x['ext']))

def find_handler(src: str, tgt: str) -> Optional[Rule]:
    src, tgt = canon(src), canon(tgt)
    for rule in RULES:
        if src in rule.sources and tgt in rule.targets and rule.check():
            return rule
    return None

# ═══════════════════════════════════════════════════════════════
# ENDPOINTS
# ═══════════════════════════════════════════════════════════════

@app.route('/api/status')
def status():
    return jsonify({
        'online': True,
        'deps': DEPS,
        'rules_available': [r.name for r in RULES if r.check()],
        'rules_disabled':  [{'name':r.name,'note':r.note} for r in RULES if not r.check()],
    })


@app.route('/api/detect', methods=['POST'])
def detect():
    f = request.files.get('file')
    if not f: return jsonify({'error':'No file'}), 400
    ext = canon(Path(f.filename).suffix)
    if ext not in FORMATS:
        return jsonify({'error':f'Unsupported format: .{ext}','filename':f.filename}), 400

    f.stream.seek(0, 2); size = f.stream.tell(); f.stream.seek(0)
    info = FORMATS[ext]
    targets = valid_targets_for(ext)
    by_cat = {}
    for t in targets: by_cat.setdefault(t['cat'], []).append(t)
    return jsonify({
        'filename': f.filename,
        'source_ext': ext,
        'source_desc': info['desc'],
        'category': info['cat'],
        'size_mb': round(size / 1048576, 2),
        'targets': targets,
        'targets_by_category': by_cat,
    })


@app.route('/api/convert', methods=['POST'])
def convert():
    f = request.files.get('file')
    target = (request.form.get('target') or '').lower().lstrip('.')
    quality = int(request.form.get('quality', 85))
    if not f: return jsonify({'error':'No file'}), 400
    if not target: return jsonify({'error':'No target'}), 400

    src = canon(Path(f.filename).suffix)
    tgt = canon(target)
    if src not in FORMATS: return jsonify({'error':f'Unsupported source: .{src}'}), 400
    if tgt not in FORMATS: return jsonify({'error':f'Unsupported target: .{tgt}'}), 400

    rule = find_handler(src, tgt)
    if rule is None:
        return jsonify({'error':f'No converter available for .{src} → .{tgt}'}), 400

    in_path = out_path = None
    try:
        in_tmp = tempfile.NamedTemporaryFile(suffix='.'+src, delete=False)
        in_path = in_tmp.name; in_tmp.close()
        f.save(in_path)
        out_path = tempfile.NamedTemporaryFile(suffix='.'+tgt, delete=False).name
        rule.handler(in_path, out_path, src, tgt, quality)
        base = Path(f.filename).stem
        return send_file(out_path, as_attachment=True,
                         download_name=f'{base}_forged.{tgt}')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if in_path and os.path.exists(in_path):
            try: os.remove(in_path)
            except: pass


if __name__ == '__main__':
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 8916
    print(f'⚒️  Dragon Forge Server on http://127.0.0.1:{port}')
    print(f'   Lanes active: {sum(1 for r in RULES if r.check())}/{len(RULES)}')
    for k, v in DEPS.items():
        print(f'     {"✓" if v else "✗"} {k}')
    app.run(host='127.0.0.1', port=port, threaded=True)
